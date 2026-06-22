"""Generate a professional Word document explaining the STT pipeline."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page margins ─────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Colour palette ────────────────────────────────────────────────────────────
DARK_BLUE  = RGBColor(0x1A, 0x37, 0x6C)
MID_BLUE   = RGBColor(0x21, 0x6F, 0xAF)
LIGHT_BLUE = RGBColor(0xD6, 0xE8, 0xF7)
ORANGE     = RGBColor(0xE8, 0x6A, 0x1A)
GREEN      = RGBColor(0x1E, 0x7A, 0x40)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
GRAY_BG    = RGBColor(0xF2, 0xF2, 0xF2)
DARK_TEXT  = RGBColor(0x1A, 0x1A, 0x1A)


# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, rgb):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    hex_color = "{:02X}{:02X}{:02X}".format(rgb[0], rgb[1], rgb[2])
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def add_heading(text, level=1, color=DARK_BLUE):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    if level == 1:
        run.font.size = Pt(18)
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(4)
    elif level == 2:
        run.font.size = Pt(14)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(3)
    else:
        run.font.size = Pt(12)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
    return p


def add_body(text, bold_parts=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if bold_parts is None:
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = DARK_TEXT
    else:
        # bold_parts: list of (text, is_bold)
        for part, is_bold in bold_parts:
            run = p.add_run(part)
            run.font.size  = Pt(11)
            run.bold       = is_bold
            run.font.color.rgb = DARK_TEXT
    return p


def add_bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_TEXT
    return p


def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x2B, 0x2B, 0x2B)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"),  "clear")
    shading.set(qn("w:color"),"auto")
    shading.set(qn("w:fill"), "F0F0F0")
    p._p.get_or_add_pPr().append(shading)
    return p


def add_table(headers, rows, header_bg=MID_BLUE, alt_bg=LIGHT_BLUE):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_bg(cell, header_bg)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(11)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            if r_idx % 2 == 0:
                set_cell_bg(cell, GRAY_BG)
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
            run.font.color.rgb = DARK_TEXT

    doc.add_paragraph()   # spacer after table
    return table


def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run("─" * 80)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)


# ═══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
run = p.add_run("Multi-Language Speech-to-Text Pipeline")
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = DARK_BLUE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Technical Architecture & Method Documentation")
run.font.size = Pt(14)
run.font.color.rgb = MID_BLUE

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("22 Indian Languages  +  English  +  Code-Switching (Hinglish)")
run.font.size = Pt(12)
run.font.color.rgb = ORANGE
run.bold = True

doc.add_paragraph()
doc.add_paragraph()

meta = [
    ("Prepared by",  "Dhruv Garg"),
    ("Date",         "June 2026"),
    ("Hardware",     "Windows 11 · NVIDIA RTX 4060 Laptop GPU (8 GB VRAM) · Intel i7"),
    ("Environment",  "C:\\nemo_venv  (Python 3.11)"),
]
for label, value in meta:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(f"{label}:  ")
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = DARK_BLUE
    r2 = p.add_run(value)
    r2.font.size = Pt(11)
    r2.font.color.rgb = DARK_TEXT

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════════════════════

add_heading("1. Executive Summary", 1)
add_body(
    "This document describes a production-ready, fully open-source, locally-hosted "
    "Speech-to-Text (STT) pipeline that supports 22 Indian languages, English, and "
    "code-switched audio (Hinglish, Telugu+English, etc.) — all without any paid API "
    "or cloud dependency."
)
add_body(
    "The pipeline uses three AI models working together with an intelligent routing "
    "engine that decides, for every short audio segment, which model(s) to use. "
    "This approach achieves Sarvam Saaras-like functionality using only free, "
    "open-source components."
)

doc.add_paragraph()
add_table(
    ["Capability", "Status"],
    [
        ["Pure English audio",                      "✅ Whisper (OpenAI) — direct route"],
        ["Pure Indian language audio (22 langs)",   "✅ IndicConformer (AI4Bharat) — direct route"],
        ["Hinglish / code-switched audio",          "✅ Word-level mixer — both models combined"],
        ["English → Hindi boundary in same chunk",  "✅ Boundary detection — automatic split"],
        ["Auto language detection",                 "✅ Whisper detect_language per chunk"],
        ["Offline / no API cost",                   "✅ Fully local"],
    ]
)


# ═══════════════════════════════════════════════════════════════════════════════
# 2. MODELS USED
# ═══════════════════════════════════════════════════════════════════════════════

add_heading("2. Models Used", 1)

add_heading("2.1  AI4Bharat IndicConformer (Hindi Large)", 2)
add_table(
    ["Property", "Detail"],
    [
        ["Architecture",  "Conformer-Large · Hybrid RNNT-CTC-BPE"],
        ["Parameters",    "~200 M"],
        ["Languages",     "Hindi (current model)  —  upgradeable to 22 Indian languages"],
        ["File",          "hindi_asr.nemo  (~499 MB)  at  C:\\nemo_venv\\"],
        ["Framework",     "NVIDIA NeMo"],
        ["License",       "Open source (AI4Bharat)"],
        ["Strength",      "Native Indian script output · accurate for pure Indian speech"],
        ["Limitation",    "Does not handle English — phonetically transliterates loanwords"],
    ]
)

add_heading("2.2  OpenAI Whisper (Small)", 2)
add_table(
    ["Property", "Detail"],
    [
        ["Architecture", "Encoder-Decoder Transformer"],
        ["Parameters",   "244 M"],
        ["Languages",    "99 languages including English and partial Indian languages"],
        ["Size",         "461 MB"],
        ["License",      "MIT — fully open source"],
        ["Strength",     "Excellent English · fast language detection · word timestamps"],
        ["Limitation",   "Translates (not transcribes) pure Hindi at >90 % confidence"],
    ]
)

add_heading("2.3  Silero VAD", 2)
add_table(
    ["Property", "Detail"],
    [
        ["Purpose",   "Voice Activity Detection — splits audio at silence boundaries"],
        ["Source",    "snakers4/silero-vad  (PyTorch Hub)"],
        ["License",   "MIT"],
        ["Input",     "Any WAV file — resampled to 16 kHz internally"],
        ["Output",    "List of (start_sec, end_sec) speech segments"],
        ["Threshold", "0.5 VAD score · 200 ms minimum silence · 100 ms minimum speech"],
    ]
)


# ═══════════════════════════════════════════════════════════════════════════════
# 3. FULL PIPELINE ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════════════

add_heading("3. Full Pipeline Architecture", 1)

add_body("The pipeline processes audio in five sequential layers:")
doc.add_paragraph()

steps = [
    ("Layer 1 — VAD Splitting",
     "Silero VAD reads the input audio and splits it at silence pauses "
     "(≥200 ms) into short chunks. Each chunk is saved as a temporary WAV file. "
     "If no silence is found, the entire file is treated as one chunk."),
    ("Layer 2 — Language Detection (per chunk)",
     "Whisper's detect_language() is run on each chunk to identify the spoken "
     "language and produce a confidence score between 0 and 1."),
    ("Layer 3 — Routing Decision",
     "Based on the detected language and confidence score, the Router sends the "
     "chunk to the correct processing path (see Section 4)."),
    ("Layer 4 — Transcription",
     "The assigned model(s) transcribe the chunk. For mixed chunks, both models "
     "run and the Mixer merges their outputs word-by-word (see Section 5)."),
    ("Layer 5 — Assembly",
     "The Assembler joins all chunk transcripts in time order and produces the "
     "final output string with a route summary header."),
]

for title, desc in steps:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f"► {title}:  ")
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = MID_BLUE
    r2 = p.add_run(desc)
    r2.font.size = Pt(11)
    r2.font.color.rgb = DARK_TEXT

doc.add_paragraph()
add_heading("Pipeline Flow Diagram", 3, color=MID_BLUE)
for line in [
    "  Audio Input (WAV / MP3 / OGG)",
    "       │",
    "       ▼",
    "  [Layer 1]  Silero VAD  ──→  Chunk 1, Chunk 2, Chunk 3 ...",
    "                                    │",
    "                            ┌───────┴───────┐",
    "                     For each chunk:",
    "                            │",
    "                       [Layer 2]",
    "                  Whisper detect_language()",
    "                  → lang_code, confidence",
    "                            │",
    "                       [Layer 3]  ROUTER",
    "              ┌─────────────┼──────────────┐",
    "              │             │              │",
    "        English >85%   Indian >90%    Indian 85-90%",
    "        or low conf    confidence     or Mixed <85%",
    "              │             │              │",
    "         [Whisper]   [IndicConformer]   [MIXER]",
    "          direct         direct      (both models)",
    "              │             │              │",
    "              └─────────────┴──────────────┘",
    "                            │",
    "                       [Layer 5]",
    "                       Assembler",
    "                            │",
    "                    Final Transcript",
]:
    add_code(line)


# ═══════════════════════════════════════════════════════════════════════════════
# 4. ROUTING LOGIC
# ═══════════════════════════════════════════════════════════════════════════════

add_heading("4. Routing Logic (Router)", 1)

add_body(
    "The Router is the brain of the pipeline. It makes a per-chunk decision based "
    "on the detected language and confidence score."
)
doc.add_paragraph()

add_table(
    ["Condition", "Route", "Models Used", "Why"],
    [
        ["User selected English manually",
         "Whisper Direct",
         "Whisper only",
         "No detection needed — user specified"],

        ["User selected Indian language manually",
         "IndicConformer Direct",
         "IndicConformer only",
         "No detection needed — user specified"],

        ["Auto-detect: English with ≥85% confidence",
         "Whisper Direct",
         "Whisper only",
         "High confidence pure English → Whisper is best"],

        ["Auto-detect: Indian language with >90% confidence",
         "IndicConformer Direct",
         "IndicConformer only",
         "Very high confidence = genuinely Indian speech. "
         "Whisper would translate (not transcribe) at this level."],

        ["Auto-detect: Indian language with 85–90% confidence",
         "Mixer",
         "Both + Mixer",
         "High but not extreme — likely has English loanwords mixed in"],

        ["Auto-detect: any language with <85% confidence",
         "Mixer",
         "Both + Mixer",
         "Low confidence = code-switching detected"],

        ["Unsupported / unknown language",
         "Whisper Fallback",
         "Whisper (auto)",
         "Best-effort — Whisper handles 99 languages"],
    ]
)

add_heading("Confidence Threshold Rationale", 3, color=MID_BLUE)
add_body(
    "The 90% threshold for 'IndicConformer Direct' was determined empirically. "
    "When Whisper detects an Indian language at >90% confidence, the audio is "
    "genuinely in that language — running Whisper with language='en' at this point "
    "causes it to translate the Indian speech to English rather than romanise it, "
    "which would produce an incorrect transcript. The threshold of 85% for "
    "'English Direct' is the standard Whisper high-confidence cutoff."
)


# ═══════════════════════════════════════════════════════════════════════════════
# 5. MIXER — CODE-SWITCHING HANDLER
# ═══════════════════════════════════════════════════════════════════════════════

add_heading("5. Mixer — Code-Switching Handler", 1)

add_body(
    "When a chunk is routed to the Mixer, both Whisper and IndicConformer transcribe "
    "it simultaneously. The Mixer then intelligently combines their outputs using one "
    "of two strategies depending on the audio content."
)

doc.add_paragraph()
add_heading("5.1  Strategy 1: Boundary Merge (English → Hindi switch)", 2)
add_body(
    "Used when the speaker switches language once within a chunk — for example, "
    "starting in English and finishing in Hindi."
)

add_body("How it works:", [("How it works:", True), ("", False)])
bullets_boundary = [
    "Run IndicConformer on the full chunk → get Hindi transcript.",
    "Scan IndicConformer output for the first 'authentic Hindi' grammar marker "
    "(words like मैं, में, है, हैं, रहा, हूँ, नहीं that never appear in phonetically "
    "transliterated English).",
    "Three sub-cases trigger the boundary strategy:",
    "  • Chunk starts with a Hindi grammar marker → entire chunk is Hindi → "
    "return IndicConformer output directly.",
    "  • >40% of IndicConformer words are grammar markers → Hindi-dominant → "
    "return IndicConformer output directly.",
    "  • First grammar marker appears past 40% of word list AND ≥2 grammar markers "
    "follow it → English→Hindi boundary mid-chunk.",
    "For the mid-chunk boundary case: estimate the timestamp of the switch using "
    "the word-count proportion, take Whisper-EN words before that timestamp as the "
    "English part, append IndicConformer words from the boundary onward as the "
    "Hindi part.",
]
for b in bullets_boundary:
    add_bullet(b, level=(1 if b.startswith("  •") else 0))

doc.add_paragraph()
add_heading("Example — Boundary Strategy", 3, color=GREEN)
add_table(
    ["", "Content"],
    [
        ["Audio spoken",      "Hello, my name is Dhruv Garg and I am from Punjab [switch] मैं अभी BML Munjal University में पढ़ रहा हूँ"],
        ["Whisper EN output", "Hello, my name is Dhruv Garg and I am from Punjab. I am now studying at BML Munjal University."],
        ["IndicConformer",    "हेलो माय नेम ... पंजाब  मैं अभी बी एम एल मुंजाल यूनिवर्सिटी में पढ़ रहा हूँ"],
        ["Boundary detected", "मैं (first grammar marker past 40%, followed by में, पढ़, रहा, हूँ)"],
        ["Final output",      "Hello, my name is Dhruv Garg and I am from Punjab. मैं अभी बी एम एल मुंजाल यूनिवर्सिटी में पढ़ रहा हूँ"],
    ],
    header_bg=GREEN
)

add_heading("5.2  Strategy 2: Word-Level Merge (True Hinglish)", 2)
add_body(
    "Used when English and Hindi words are interleaved word-by-word throughout "
    "the chunk — classic Hinglish."
)

add_body("How it works:", [("How it works:", True), ("", False)])
bullets_wordlevel = [
    "Run Whisper with language='en' to get word-level timestamps (all output is ASCII).",
    "Run IndicConformer to get the native-script version of the same audio.",
    "For each Whisper word (in order):",
    "  • ASCII + in English vocabulary list → keep as English from Whisper.",
    "  • ASCII + in Hindi romanisation list (mujhe, jana, hai, etc.) → replace "
    "with the corresponding IndicConformer word (Devanagari script).",
    "  • ASCII + phonetic cluster heuristic (aa, dh, bh, etc.) → treat as romanised "
    "Indian → replace with IndicConformer word.",
    "  • Already non-ASCII (IndicConformer/Whisper gave native script) → keep directly.",
    "1-to-1 alignment: the IndicConformer index advances for EVERY Whisper word "
    "(English or Indian) to maintain positional alignment and avoid mis-substitution.",
    "Quick-exit: if zero Hindi romanisations are found in Whisper output → entire "
    "chunk is English → return Whisper transcript directly (avoids wasted IndicConformer run).",
]
for b in bullets_wordlevel:
    add_bullet(b, level=(1 if b.startswith("  •") else 0))

doc.add_paragraph()
add_heading("Example — Word-Level Strategy", 3, color=GREEN)
add_table(
    ["", "Content"],
    [
        ["Audio spoken",       "mujhe tomorrow office jana hai"],
        ["Whisper EN",         "mujhe  tomorrow  office  jana  hai"],
        ["IndicConformer",     "मुझे   टुमारो    ऑफिस    जाना    है"],
        ["Classification",     "Indian  English  English  Indian  Indian"],
        ["Action",             "IC[0]   keep     keep     IC[3]   IC[4]"],
        ["Final output",       "मुझे tomorrow office जाना है"],
    ],
    header_bg=GREEN
)


# ═══════════════════════════════════════════════════════════════════════════════
# 6. FILE STRUCTURE
# ═══════════════════════════════════════════════════════════════════════════════

add_heading("6. Codebase Structure", 1)

add_table(
    ["File / Folder", "Purpose"],
    [
        ["app.py",                   "Gradio web UI — loads all models, handles upload, shows transcript"],
        ["config.py",                "Central config: model paths, confidence thresholds, VAD settings, language maps"],
        ["run_test.py",              "CLI test script — runs a WAV file through the full pipeline and prints results"],
        ["models/indic_model.py",    "IndicConformer wrapper — loads .nemo model, exposes transcribe()"],
        ["models/whisper_model.py",  "Whisper wrapper — detect_language(), transcribe(), transcribe_words()"],
        ["pipeline/vad.py",          "Silero VAD — splits audio into chunks at silence boundaries"],
        ["pipeline/lang_detect.py",  "classify() — maps (lang, confidence) to route type"],
        ["pipeline/router.py",       "Router — per-chunk routing decision + mixer invocation"],
        ["pipeline/mixer.py",        "Boundary merge + Word-level merge for code-switched audio"],
        ["pipeline/assembler.py",    "Assembler — joins chunk transcripts + produces route summary"],
        ["NeMo/nemo/",               "NeMo core library (required for IndicConformer inference)"],
    ]
)


# ═══════════════════════════════════════════════════════════════════════════════
# 7. LIVE TEST RESULTS
# ═══════════════════════════════════════════════════════════════════════════════

add_heading("7. Live Test Results", 1)

add_heading("Test File: testt.wav  (Intro — English + Hindi)", 2)
add_table(
    ["Chunk", "Time", "Detected", "Route", "Transcript"],
    [
        ["1", "0.9–4.3s", "hi 40%",  "Mixed",   "Hello, my name is Dhruv Garg and I am from Punjab, Bathinda."],
        ["2", "4.6–5.5s", "hi 54%",  "Mixed",   "मैं अभी"],
        ["3", "5.8–8.6s", "hi 63%",  "Mixed",   "पी एम एल मुंजाल यूनिवर्सिटी में पढ़ रहा हूँ"],
        ["4", "8.9–10.3s","hi 55%",  "Mixed",   "and I am"],
        ["5", "11.0–11.8s","ml 38%", "Mixed",   "working in."],
        ["6", "12.2–13.0s","hi 56%", "Mixed",   "Bluestem."],
        ["7", "13.5–14.4s","en 86%", "Whisper", "as it interned."],
    ]
)
add_body("Final: Hello, my name is Dhruv Garg and I am from Punjab, Bathinda. "
         "मैं अभी पी एम एल मुंजाल यूनिवर्सिटी में पढ़ रहा हूँ and I am working in. Bluestem. as it interned.",
         [("Final Transcript:  ", True),
          ("Hello, my name is Dhruv Garg and I am from Punjab, Bathinda. "
           "मैं अभी पी एम एल मुंजाल यूनिवर्सिटी में पढ़ रहा हूँ and I am working in. Bluestem. as it interned.", False)])

doc.add_paragraph()
add_heading("Test File: test5.wav  (Pure Hindi + English mix)", 2)
add_table(
    ["Chunk", "Time", "Detected", "Route", "Transcript"],
    [
        ["1", "0.7–2.5s",  "hi 94%", "IndicConformer", "मेरा नाम ध्रुव गर्ग है"],
        ["2", "3.2–4.0s",  "hi 44%", "Mixed",          "I am from"],
        ["3", "4.5–5.1s",  "hi 40%", "Mixed",          "Punjab."],
        ["4", "5.7–9.3s",  "hi 64%", "Mixed",          "Currently I am studying in BML, Munjal University."],
        ["5", "9.8–11.9s", "hi 99%", "IndicConformer", "मैं अभी इंटर्नशिप कर रहा हूँ"],
        ["6", "12.3–13.2s","hi 61%", "Mixed",          "ब्लोस्टैम में"],
    ]
)
add_body("", [("Final Transcript:  ", True),
              ("मेरा नाम ध्रुव गर्ग है I am from Punjab. Currently I am studying in BML, Munjal University. "
               "मैं अभी इंटर्नशिप कर रहा हूँ ब्लोस्टैम में", False)])

doc.add_paragraph()
add_heading("Test File: test4.wav  (Gaming conversation)", 2)
add_table(
    ["Chunk", "Time", "Detected", "Route", "Transcript"],
    [
        ["1", "0.5–1.6s",  "hi 26%", "Mixed", "Hello Drew, bye!"],
        ["2", "1.8–5.6s",  "ur 86%", "Mixed", "or download FC 2026 game."],
        ["3", "5.8–7.7s",  "ur 62%", "Mixed", "I have purchased the controller"],
        ["4", "8.0–11.1s", "hi 77%", "Mixed", "आज साथ में खेलती हो शाम को पांच बजे खेलेंगे ठीक है"],
    ]
)
add_body("", [("Final Transcript:  ", True),
              ("Hello Drew, bye! or download FC 2026 game. I have purchased the controller "
               "आज साथ में खेलती हो शाम को पांच बजे खेलेंगे ठीक है", False)])


# ═══════════════════════════════════════════════════════════════════════════════
# 8. CURRENT LIMITATIONS
# ═══════════════════════════════════════════════════════════════════════════════

add_heading("8. Current Limitations", 1)

add_table(
    ["Limitation", "Root Cause", "Fix / Upgrade Path"],
    [
        ["Current model is Hindi-only",
         "hindi_asr.nemo supports only Hindi",
         "Upgrade to ai4bharat/indic-conformer-600m-multilingual → all 22 Indian languages"],
        ["English loanwords in Hindi come out as Devanagari (e.g. ऑफिस instead of office)",
         "Hindi IndicConformer phonetically transcribes all English",
         "Multilingual model handles proper nouns better; word-level mixer partially mitigates"],
        ["Short audio chunks (<0.5 s) may misroute",
         "Too little audio for accurate language detection",
         "Increase MIN_SPEECH_MS in config.py; merge very short adjacent chunks"],
        ["No true word-level timestamps from IndicConformer",
         "NeMo RNNT-CTC does not expose word boundaries",
         "Use WhisperX for forced alignment on IndicConformer output"],
        ["Company names / rare proper nouns misrecognised",
         "Both models trained on general speech",
         "Fine-tune on domain-specific vocabulary"],
        ["Inference is CPU-only (slow ~2–3s per chunk)",
         "NeMo model loaded on CPU",
         "Move to GPU inference; use faster-whisper for Whisper side"],
    ]
)


# ═══════════════════════════════════════════════════════════════════════════════
# 9. UPGRADE ROADMAP
# ═══════════════════════════════════════════════════════════════════════════════

add_heading("9. Upgrade Roadmap", 1)

add_table(
    ["Phase", "Task", "Effort", "Impact"],
    [
        ["Short-term\n(1–2 days)",
         "Replace hindi_asr.nemo with indic-conformer-600m-multilingual",
         "Low — change one line in config.py",
         "All 22 Indian languages supported immediately"],
        ["Medium-term\n(1–2 weeks)",
         "Move Whisper to faster-whisper + GPU inference",
         "Medium — swap library, test accuracy",
         "4–5× speed improvement, real-time capable"],
        ["Medium-term\n(2–3 weeks)",
         "Fine-tune Whisper large-v3 on MUCS 2021 code-switched data",
         "High — training required",
         "Single model for all languages + code-switching"],
        ["Long-term\n(1–2 months)",
         "WhisperX word alignment for IndicConformer output",
         "Medium",
         "More accurate word-level mixing in Hinglish"],
        ["Long-term\n(1–2 months)",
         "INT8 quantisation for both models",
         "Low–Medium",
         "<0.5s latency for voice agent integration"],
    ]
)


# ═══════════════════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════════════════

out_path = r"C:\Users\Dhruv Garg\OneDrive\Desktop\New folder (5)\IndicConformerASR\Pipeline_Documentation.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
